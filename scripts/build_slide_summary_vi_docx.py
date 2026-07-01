# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
import math
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_professional_report_en import (
    COLORS,
    LOGO_PATH,
    METRICS_CSV,
    METRICS_JSON,
    SCREENSHOT_DIR,
    draw_rounded_rectangle,
    load_font,
    pil_color,
    rgb_color,
    set_run_font,
    set_style_font,
    set_table_borders,
    set_table_cell_margins,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "slide_summary_vi_professional.docx"
ASSET_DIR = ROOT / "reports" / "slide_summary_vi_assets"

SLIDE_W = 10.0


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, size: int, fill: str, bold: bool = False, anchor: str | None = None) -> None:
    draw.text(xy, text, font=load_font(size, bold=bold), fill=pil_color(fill), anchor=anchor)


def load_evidence() -> tuple[dict, dict[str, str]]:
    with METRICS_JSON.open("r", encoding="utf-8") as handle:
        evidence = json.load(handle)
    metrics: dict[str, str] = {}
    with METRICS_CSV.open("r", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            metrics[row["metric"]] = row["value"]
    return evidence, metrics


def metric(metrics: dict[str, str], key: str, default: float = 0.0) -> float:
    try:
        return float(metrics.get(key, default))
    except (TypeError, ValueError):
        return default


def pct(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%"


def create_architecture_vi(out_path: Path) -> None:
    width, height = 1700, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 48), "Luồng kiến trúc MLOps", 52, COLORS["navy"], True)
    draw_text(draw, (70, 112), "Từ dữ liệu server metrics đến phát hiện anomaly có giám sát, audit và retraining.", 27, COLORS["gray"])

    stages = [
        ("Raw server metrics", "CPU, memory,\nrequest count,\nerror rate, latency"),
        ("Airflow pipeline", "Validation,\npreprocessing,\nfeature engineering"),
        ("Isolation Forest", "Huấn luyện model\nvà promotion gates"),
        ("MLflow registry", "Tracking, artifacts,\nproduction version"),
        ("FastAPI serving", "/detect, /explain,\n/drift, /retrain"),
        ("PostgreSQL audit", "Prediction logs,\nrequest_id,\nmodel version"),
        ("Prometheus/Grafana", "Latency,\nanomaly rate,\nerrors, drift score"),
        ("Alertmanager bridge", "Cảnh báo kích hoạt\nAirflow retraining"),
    ]
    coords = [
        (80, 205),
        (465, 205),
        (850, 205),
        (1235, 205),
        (80, 520),
        (465, 520),
        (850, 520),
        (1235, 520),
    ]
    box_w, box_h = 315, 160
    accents = [COLORS["blue"], COLORS["orange"], COLORS["green"], COLORS["navy"]]

    for idx, ((title, body), (x, y)) in enumerate(zip(stages, coords)):
        accent = accents[idx % len(accents)]
        draw_rounded_rectangle(draw, (x, y, x + box_w, y + box_h), 18, "FFFFFF", "D9E2EC", 2)
        draw.rectangle((x, y, x + box_w, y + 12), fill=pil_color(accent))
        draw_text(draw, (x + 24, y + 36), title, 26, COLORS["navy"], True)
        for line_idx, line in enumerate(body.split("\n")):
            draw_text(draw, (x + 24, y + 78 + line_idx * 30), line, 22, COLORS["gray"])

    def arrow(start: tuple[int, int], end: tuple[int, int], color: str = COLORS["gray"]) -> None:
        draw.line((start, end), fill=pil_color(color), width=4)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        size = 14
        p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
        p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
        draw.polygon([end, p1, p2], fill=pil_color(color))

    def polyline(points: list[tuple[int, int]], color: str) -> None:
        for start, end in zip(points, points[1:]):
            draw.line((start, end), fill=pil_color(color), width=4)
        start, end = points[-2], points[-1]
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        size = 14
        p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
        p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
        draw.polygon([end, p1, p2], fill=pil_color(color))

    for i in range(3):
        arrow((coords[i][0] + box_w, coords[i][1] + box_h // 2), (coords[i + 1][0] - 20, coords[i + 1][1] + box_h // 2))
    arrow((coords[3][0] + box_w // 2, coords[3][1] + box_h), (coords[7][0] + box_w // 2, coords[7][1] - 22))
    arrow((coords[4][0] + box_w, coords[4][1] + box_h // 2), (coords[5][0] - 20, coords[5][1] + box_h // 2))
    arrow((coords[5][0] + box_w, coords[5][1] + box_h // 2), (coords[6][0] - 20, coords[6][1] + box_h // 2))
    arrow((coords[6][0] + box_w, coords[6][1] + box_h // 2), (coords[7][0] - 20, coords[7][1] + box_h // 2))
    arrow((coords[1][0] + box_w // 2, coords[1][1] + box_h), (coords[5][0] + box_w // 2, coords[5][1] - 22), COLORS["blue"])

    # Feedback path runs outside the component boxes to avoid crossing labels.
    polyline(
        [
            (coords[7][0] + box_w // 2, coords[7][1] + box_h),
            (coords[7][0] + box_w // 2, 792),
            (40, 792),
            (40, coords[1][1] + box_h + 28),
            (coords[1][0] - 8, coords[1][1] + box_h + 28),
        ],
        COLORS["orange"],
    )
    draw_text(draw, (1110, 800), "feedback retraining", 22, COLORS["orange"], True)
    draw_text(draw, (70, 850), "Stack evidence: Airflow, MLflow, FastAPI, PostgreSQL, Prometheus, Grafana, Alertmanager.", 20, COLORS["gray"])
    img.save(out_path, quality=95)


def create_kpi_vi(metrics: dict[str, str], out_path: Path) -> None:
    width, height = 1700, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "Bảng số liệu chính", 52, COLORS["navy"], True)
    draw_text(draw, (70, 112), "Các chỉ số được lấy từ reports/demo_evidence và bằng chứng runtime.", 26, COLORS["gray"])
    cards = [
        ("Precision", pct(metric(metrics, "model_precision")), COLORS["blue"], "Độ đúng khi dự đoán anomaly"),
        ("Recall", pct(metric(metrics, "model_recall")), COLORS["green"], "Tỷ lệ anomaly được phát hiện"),
        ("F1-score", pct(metric(metrics, "model_f1_score")), COLORS["orange"], "Cân bằng precision/recall"),
        ("API p95", f"{metric(metrics, 'api_p95_latency_seconds'):.2f}s", COLORS["navy"], "Độ trễ p95 của API"),
    ]
    x, y = 70, 190
    card_w, card_h, gap = 370, 170, 35
    for idx, (label, value, accent, note) in enumerate(cards):
        left = x + idx * (card_w + gap)
        draw_rounded_rectangle(draw, (left, y, left + card_w, y + card_h), 22, "FFFFFF", "D9E2EC", 2)
        draw.rectangle((left, y, left + 12, y + card_h), fill=pil_color(accent))
        draw_text(draw, (left + 34, y + 30), label, 27, COLORS["gray"], True)
        draw_text(draw, (left + 34, y + 78), value, 50, accent, True)
        draw_text(draw, (left + 34, y + 135), note, 20, COLORS["gray"])

    bars = [
        ("Production requests", metric(metrics, "production_requests"), 400, COLORS["blue"], "400"),
        ("Production anomalies", metric(metrics, "production_anomalies"), 400, COLORS["orange"], "108"),
        ("Load-test RPS", metric(metrics, "loadtest_rps"), 10, COLORS["green"], f"{metric(metrics, 'loadtest_rps'):.2f}"),
        ("Load-test p95", metric(metrics, "loadtest_p95_ms"), 500, COLORS["navy"], "210 ms"),
        ("Failure rate", metric(metrics, "loadtest_failure_rate"), 1, "9B1C1C", "0"),
    ]
    base_y = 440
    for idx, (label, value, maximum, accent, shown) in enumerate(bars):
        row_y = base_y + idx * 48
        draw_text(draw, (90, row_y), label, 23, COLORS["black"], True)
        x_bar, y_bar, w_bar, h_bar = 500, row_y + 5, 820, 24
        draw_rounded_rectangle(draw, (x_bar, y_bar, x_bar + w_bar, y_bar + h_bar), 12, "EEF2F6")
        fill_w = int(min(w_bar, max(0, value / maximum * w_bar))) if maximum else 0
        if fill_w:
            draw_rounded_rectangle(draw, (x_bar, y_bar, x_bar + fill_w, y_bar + h_bar), 12, accent)
        draw_text(draw, (x_bar + w_bar + 35, row_y), shown, 23, accent, True)
    img.save(out_path, quality=95)


def create_feature_vi(evidence: dict, out_path: Path) -> None:
    features = evidence["api"]["explain"]["body"]["top_features"][:8]
    width, height = 1600, 740
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "Explainability: yếu tố làm tăng rủi ro anomaly", 44, COLORS["navy"], True)
    draw_text(draw, (70, 105), "Phương pháp: SHAP tree explainer. Giá trị càng cao, ảnh hưởng càng mạnh.", 24, COLORS["gray"])
    max_val = max(item["abs_impact"] for item in features)
    x_label, x_bar, bar_w = 90, 560, 760
    for idx, item in enumerate(features):
        y = 180 + idx * 62
        label = item["feature"].replace("_", " ")
        draw_text(draw, (x_label, y), label, 23, COLORS["black"], True)
        draw_rounded_rectangle(draw, (x_bar, y + 6, x_bar + bar_w, y + 34), 14, "EEF2F6")
        fill_w = int(item["abs_impact"] / max_val * bar_w)
        accent = COLORS["orange"] if idx < 3 else COLORS["blue"]
        draw_rounded_rectangle(draw, (x_bar, y + 6, x_bar + fill_w, y + 34), 14, accent)
        draw_text(draw, (x_bar + bar_w + 30, y + 2), f"{item['abs_impact']:.3f}", 23, accent, True)
    img.save(out_path, quality=95)


def create_fairness_vi(evidence: dict, out_path: Path) -> None:
    groups = evidence["api"]["fairness"]["body"]["group_metrics"]
    width, height = 1600, 700
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "Fairness vận hành theo server_id", 44, COLORS["navy"], True)
    draw_text(draw, (70, 105), "Không dùng thuộc tính con người; kiểm tra chênh lệch anomaly rate theo phân đoạn vận hành.", 24, COLORS["gray"])
    x_label, x_bar, bar_w = 100, 470, 780
    for idx, (group, values) in enumerate(groups.items()):
        y = 185 + idx * 78
        rate = values["anomaly_rate"]
        draw_text(draw, (x_label, y), group, 26, COLORS["black"], True)
        draw_text(draw, (x_label, y + 34), f"{values['request_count']} requests", 19, COLORS["gray"])
        draw_rounded_rectangle(draw, (x_bar, y + 12, x_bar + bar_w, y + 44), 16, "EEF2F6")
        fill_w = int(rate * bar_w)
        accent = COLORS["green"] if rate < 0.35 else COLORS["orange"]
        if fill_w:
            draw_rounded_rectangle(draw, (x_bar, y + 12, x_bar + fill_w, y + 44), 16, accent)
        draw_text(draw, (x_bar + bar_w + 30, y + 8), f"{rate:.3f}", 24, accent, True)
    img.save(out_path, quality=95)


def create_assets(evidence: dict, metrics: dict[str, str]) -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "architecture": ASSET_DIR / "architecture_vi_fixed.png",
        "kpi": ASSET_DIR / "kpi_vi.png",
        "feature": ASSET_DIR / "feature_vi.png",
        "fairness": ASSET_DIR / "fairness_vi.png",
    }
    create_architecture_vi(assets["architecture"])
    create_kpi_vi(metrics, assets["kpi"])
    create_feature_vi(evidence, assets["feature"])
    create_fairness_vi(evidence, assets["fairness"])
    return assets


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.50)
    section.right_margin = Inches(0.50)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.20)

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11, COLORS["black"])
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    set_style_font(styles["Heading 1"], "Calibri", 24, "2E74B5", True)
    styles["Heading 1"].paragraph_format.space_before = Pt(0)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 2"], "Calibri", 14, "1F4D78", True)
    styles["Heading 2"].paragraph_format.space_before = Pt(6)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Slide summary | final_ddm501")
    set_run_font(run, size=8.5, color=COLORS["gray"])

    doc.core_properties.title = "Tóm tắt slide final_ddm501"
    doc.core_properties.subject = "Bản slide summary tiếng Việt có hình ảnh và bảng số liệu"
    doc.core_properties.author = "Mai Văn An, Nguyễn Thanh Bình, Phạm Đình Hoàng, Nguyễn Tuấn Dũng"


def add_slide_title(doc: Document, number: int, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"SLIDE {number:02d}")
    set_run_font(r, size=9, color=COLORS["orange"], bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=24, color=COLORS["navy"], bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(subtitle)
        set_run_font(r, size=11.5, color=COLORS["gray"])


def add_slide_image(doc: Document, path: Path, width: float, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    alt_text = caption or f"Hình minh họa {path.stem}"
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", path.stem)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(caption)
        set_run_font(r, size=7.5, color=COLORS["gray"], italic=True)


def add_image_grid(doc: Document, items: list[tuple[Path, str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(items))
    set_table_geometry(table, [SLIDE_W / len(items)] * len(items), indent_dxa=0)
    set_table_borders(table, color="FFFFFF", size="0")
    set_table_cell_margins(table, top=20, bottom=20, start=80, end=80)
    mark_header(table.rows[0])
    for idx, (path, caption) in enumerate(items):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(path), width=Inches(widths[idx]))
        shape._inline.docPr.set("descr", caption)
        shape._inline.docPr.set("title", path.stem)
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(caption)
        set_run_font(r, size=7.3, color=COLORS["gray"], italic=True)


def mark_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def set_cell(cell, text: str, size: float = 10, color: str = COLORS["black"], bold: bool = False, fill: str | None = None, align: int | None = None) -> None:
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = "E8EEF5", font_size: float = 8.8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_table_cell_margins(table, top=80, bottom=80, start=110, end=110)
    mark_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell(table.rows[0].cells[idx], header, font_size, COLORS["navy"], True, header_fill)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value, font_size)
    set_table_geometry(table, widths)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [SLIDE_W / len(metrics)] * len(metrics)
    set_table_geometry(table, widths, indent_dxa=0)
    set_table_borders(table, color="D9E2EC", size="6")
    set_table_cell_margins(table, top=100, bottom=100, start=140, end=140)
    mark_header(table.rows[0])
    fills = ["EAF4FF", "EAF8EF", "FFF3EA", "F2F4F7"]
    for idx, (label, value, note) in enumerate(metrics):
        cell = table.rows[0].cells[idx]
        set_cell(cell, "", fill=fills[idx % len(fills)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=9.2, color=COLORS["gray"], bold=True)
        r = p.add_run(value + "\n")
        set_run_font(r, size=18, color=[COLORS["blue"], COLORS["green"], COLORS["orange"], COLORS["navy"]][idx % 4], bold=True)
        r = p.add_run(note)
        set_run_font(r, size=8.2, color=COLORS["gray"])


def add_bullets(doc: Document, items: list[str], size: float = 10.5) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=size)


def add_cover(doc: Document, evidence: dict) -> None:
    add_slide_title(doc, 1, "Tóm tắt slide dự án final_ddm501", "Production-style MLOps for Server Log Anomaly Detection")
    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        shape = run.add_picture(str(LOGO_PATH), width=Inches(4.2))
        shape._inline.docPr.set("descr", "Logo FPT School of Business & Technology")
        shape._inline.docPr.set("title", "FPT logo")
    add_metric_strip(
        doc,
        [
            ("Model", "Isolation Forest", "Anomaly detection"),
            ("Evidence", "400 requests", "Production audit"),
            ("Runtime", "0.05s p95", "FastAPI latency"),
            ("CI/CD", "Quality gates", "Tests, build, scan"),
        ],
    )
    doc.add_paragraph()
    rows = [
        ["Mai Văn An", "Project Lead và kiến trúc MLOps"],
        ["Nguyễn Thanh Bình", "API serving, security, runtime validation"],
        ["Phạm Đình Hoàng", "Model training, evaluation, explainability"],
        ["Nguyễn Tuấn Dũng", "Observability, CI/CD, documentation"],
    ]
    add_data_table(doc, ["Thành viên", "Vai trò chính trong báo cáo"], rows, [3.0, 7.0], "FFF3EA", 10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Thời điểm thu thập evidence: {evidence.get('collected_at', '2026-07-01T20:52:31+07:00')}")
    set_run_font(r, size=9.5, color=COLORS["gray"], italic=True)
    doc.add_page_break()


def slide_objective(doc: Document, metrics: dict[str, str], assets: dict[str, Path]) -> None:
    add_slide_title(doc, 2, "Mục tiêu và thông điệp chính", "Dự án không chỉ train model, mà đóng gói thành một stack MLOps có thể vận hành và quan sát.")
    add_metric_strip(
        doc,
        [
            ("Precision", pct(metric(metrics, "model_precision")), "Model quality"),
            ("Recall", pct(metric(metrics, "model_recall")), "Anomaly captured"),
            ("F1-score", pct(metric(metrics, "model_f1_score")), "Balanced metric"),
            ("Requests", f"{metric(metrics, 'production_requests'):.0f}", "Audit evidence"),
        ],
    )
    add_bullets(
        doc,
        [
            "Xây dựng hệ thống phát hiện anomaly trên server metrics: CPU, memory, request count, error rate, latency.",
            "Bao phủ vòng đời MLOps: data validation, feature engineering, model registry, serving, audit log, monitoring, alert, retraining.",
            "Không triển khai concept drift có label vì production chưa có ground-truth; thay vào đó theo dõi data drift và prediction distribution drift.",
            "Điểm mạnh khi thuyết trình: có evidence thật từ FastAPI, PostgreSQL, Prometheus/Grafana, MLflow, Airflow và GitHub Actions.",
        ],
        11,
    )
    add_slide_image(doc, assets["kpi"], 8.0, "Hình 1. Bảng số liệu chính từ evidence.")
    doc.add_page_break()


def slide_architecture(doc: Document, assets: dict[str, Path]) -> None:
    add_slide_title(doc, 3, "Kiến trúc MLOps", "Mũi tên màu cam là luồng feedback retraining, đã được vẽ vòng ngoài để không đè lên các component.")
    add_slide_image(doc, assets["architecture"], 8.4, "Hình 2. Luồng kiến trúc MLOps đã chỉnh mũi tên feedback.")
    rows = [
        ["Lifecycle", "Airflow + MLflow + registry", "Training, tracking, artifacts và production version."],
        ["Serving & audit", "FastAPI + PostgreSQL", "Inference endpoint, API key, prediction log và request_id."],
        ["Monitoring & retraining", "Prometheus/Grafana + Alertmanager", "Giám sát latency, anomaly rate, drift score và kích hoạt retraining."],
    ]
    add_data_table(doc, ["Nhóm", "Thành phần", "Ý nghĩa"], rows, [2.0, 3.0, 5.0], "EAF4FF", 8.0)
    doc.add_page_break()


def slide_model(doc: Document, evidence: dict, metrics: dict[str, str]) -> None:
    add_slide_title(doc, 4, "Kết quả model", "Model Isolation Forest đạt kết quả mạnh trên bộ validation/evidence của demo.")
    rows = [
        ["Precision", "0.973684", "Tỷ lệ dự đoán anomaly đúng cao."],
        ["Recall", "0.973684", "Tỷ lệ anomaly được bắt được cao."],
        ["F1-score", "0.973684", "Cân bằng tốt giữa precision và recall."],
        ["False positive rate", "0.001466", "Tỷ lệ báo nhầm thấp trong validation evidence."],
        ["Training rows", "720", "Số dòng dữ liệu sau validation."],
        ["Production model", evidence["api"]["health"]["body"]["model_version"], "FastAPI load model từ local registry."],
    ]
    add_data_table(doc, ["Chỉ số", "Giá trị", "Diễn giải"], rows, [2.0, 1.5, 6.5], "E8EEF5", 9.5)
    mlflow = SCREENSHOT_DIR / "mlflow_experiment.png"
    if mlflow.exists():
        add_slide_image(doc, mlflow, 6.2, "Hình 3. MLflow experiment evidence cho model tracking.")
    doc.add_page_break()


def slide_runtime(doc: Document, metrics: dict[str, str]) -> None:
    add_slide_title(doc, 5, "Production runtime và load test", "Serving path có audit log, metric quan sát và smoke load test không ghi nhận failure.")
    rows = [
        ["Production requests", f"{metric(metrics, 'production_requests'):.0f}", "PostgreSQL prediction_logs"],
        ["Production anomalies", f"{metric(metrics, 'production_anomalies'):.0f}", "PostgreSQL prediction_logs"],
        ["Rolling anomaly rate", "0.375", "Prometheus"],
        ["API p95 latency", f"{metric(metrics, 'api_p95_latency_seconds'):.2f}s", "Prometheus"],
        ["API error count", "0", "Prometheus"],
        ["Load-test RPS", f"{metric(metrics, 'loadtest_rps'):.2f}", "Locust CSV"],
        ["Load-test p95 / p99", "210 ms / 360 ms", "Locust CSV"],
        ["Load-test failure rate", "0", "Locust CSV"],
    ]
    add_data_table(doc, ["Runtime metric", "Giá trị", "Nguồn"], rows, [3.2, 2.0, 4.8], "EAF8EF", 8.8)
    grafana = SCREENSHOT_DIR / "grafana_dashboard.png"
    if grafana.exists():
        add_slide_image(doc, grafana, 6.0, "Hình 4. Grafana dashboard theo dõi serving và model telemetry.")
    doc.add_page_break()


def slide_observability(doc: Document) -> None:
    add_slide_title(doc, 6, "Observability và demo UI", "Dự án có đủ bề mặt quan sát để chứng minh khả năng vận hành.")
    streamlit = SCREENSHOT_DIR / "streamlit_dashboard.png"
    fastapi = SCREENSHOT_DIR / "fastapi_docs.png"
    grid_items = []
    widths = []
    if streamlit.exists():
        grid_items.append((streamlit, "Hình 5. Streamlit dashboard cho demo input và production predictions."))
        widths.append(4.55)
    if fastapi.exists():
        grid_items.append((fastapi, "Hình 6. FastAPI docs cho các endpoint phục vụ inference và vận hành."))
        widths.append(4.55)
    if grid_items:
        add_image_grid(doc, grid_items, widths)
    rows = [
        ["Prometheus", "Scrape metrics từ FastAPI: request count, error count, latency, anomaly rate, drift score."],
        ["Grafana", "Dashboard hóa request rate, p95 latency, anomaly rate, prediction count, latest drift score."],
        ["Alertmanager", "Routing cảnh báo DataDriftDetected và HighAnomalyRate sang bridge."],
    ]
    add_data_table(doc, ["Thành phần", "Vai trò trong observability"], rows, [2.3, 7.7], "F2F4F7", 9)
    doc.add_page_break()


def slide_explain_fairness(doc: Document, evidence: dict, assets: dict[str, Path]) -> None:
    add_slide_title(doc, 7, "Explainability và fairness", "Giải thích dự đoán bằng SHAP và kiểm tra chênh lệch anomaly rate theo server_id.")
    add_image_grid(
        doc,
        [
            (assets["feature"], "Hình 7. Top feature impacts làm tăng rủi ro anomaly."),
            (assets["fairness"], "Hình 8. Fairness vận hành theo server_id."),
        ],
        [4.8, 4.8],
    )
    fairness = evidence["api"]["fairness"]["body"]
    rows = [
        [group, str(values["request_count"]), f"{values['anomaly_rate']:.6f}", f"{values['high_risk_rate']:.6f}"]
        for group, values in fairness["group_metrics"].items()
    ]
    add_data_table(doc, ["Group", "Requests", "Anomaly rate", "High-risk rate"], rows, [2.3, 1.7, 2.8, 3.2], "FFF3EA", 7.6)
    doc.add_page_break()


def slide_retraining_cicd(doc: Document, evidence: dict) -> None:
    add_slide_title(doc, 8, "Retraining, CI/CD và kết luận", "Stack có đường dẫn từ drift/alert đến retraining, cộng với quality gates trong GitHub Actions.")
    drift = evidence["api"]["drift"]["body"]
    retrain = evidence["api"]["retrain_check"]["body"]
    rows = [
        ["Data drift detected", str(drift["data_drift_detected"]), "FastAPI /drift"],
        ["Prediction drift detected", str(drift["prediction_drift_detected"]), "FastAPI /drift"],
        ["Max data drift score", f"{drift['max_data_drift_score']:.6f}", "FastAPI /drift"],
        ["Retraining triggered", str(retrain["retraining_triggered"]), retrain["reason"]],
        ["Active alerts", "DataDriftDetected, HighAnomalyRate", "Prometheus ALERTS"],
    ]
    add_data_table(doc, ["Kiểm tra", "Giá trị", "Nguồn/Ghi chú"], rows, [3.0, 3.0, 4.0], "EAF4FF", 8.8)
    airflow = SCREENSHOT_DIR / "airflow_retraining_dag.png"
    if airflow.exists():
        add_slide_image(doc, airflow, 4.9, "Hình 9. Airflow retraining DAG.")
    add_bullets(
        doc,
        [
            "Kết luận: đây là demo MLOps production-style hoàn chỉnh, có model, API, audit, monitoring, alert và CI/CD.",
            "Bước tiếp theo: harden secrets, lưu persistent artifacts/registry, thêm ground-truth production labels và load test lớn hơn.",
        ],
        10,
    )


def build() -> None:
    evidence, metrics = load_evidence()
    assets = create_assets(evidence, metrics)
    doc = Document()
    configure_doc(doc)
    add_cover(doc, evidence)
    slide_objective(doc, metrics, assets)
    slide_architecture(doc, assets)
    slide_model(doc, evidence, metrics)
    slide_runtime(doc, metrics)
    slide_observability(doc)
    slide_explain_fairness(doc, evidence, assets)
    slide_retraining_cicd(doc, evidence)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
