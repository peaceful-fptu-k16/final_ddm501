# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import json
import math
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "final_ddm501_professional_report_en.docx"
ASSET_DIR = ROOT / "reports" / "professional_report_assets"
LOGO_PATH = Path(
    r"C:\Users\thanh\AppData\Local\Temp\codex-clipboard-44843ba4-f41c-4730-8ec4-0e9d5b1793b5.png"
)
METRICS_JSON = ROOT / "reports" / "demo_evidence" / "latest_metrics.json"
METRICS_CSV = ROOT / "reports" / "demo_evidence" / "metrics_table.csv"
SCREENSHOT_DIR = ROOT / "docs" / "assets" / "screenshots"


COLORS = {
    "navy": "0B2545",
    "blue": "0072BC",
    "orange": "F37021",
    "green": "00A651",
    "gray": "667085",
    "light_gray": "F2F4F7",
    "line": "D9E2EC",
    "soft_blue": "EAF4FF",
    "soft_orange": "FFF3EA",
    "soft_green": "EAF8EF",
    "white": "FFFFFF",
    "black": "111827",
}


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.strip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def pil_color(value: str | None) -> str | None:
    if value is None:
        return None
    if value.startswith("#") or value.lower() in {"white", "black"}:
        return value
    if len(value) == 6 and all(ch in "0123456789abcdefABCDEF" for ch in value):
        return f"#{value}"
    return value


def rgb_color(value: str) -> RGBColor:
    return RGBColor(*hex_to_rgb(value))


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf"),
    ]
    for candidate in font_candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_rounded_rectangle(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    radius: int,
    fill: str,
    outline: str | None = None,
    width: int = 1,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=pil_color(fill), outline=pil_color(outline), width=width)


def draw_text(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    size: int = 32,
    fill: str = COLORS["black"],
    bold: bool = False,
    anchor: str | None = None,
) -> None:
    draw.text(xy, text, font=load_font(size, bold=bold), fill=f"#{fill}", anchor=anchor)


def text_size(draw: ImageDraw.ImageDraw, text: str, size: int = 32, bold: bool = False) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=load_font(size, bold=bold))
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def load_data() -> tuple[dict, dict[str, str]]:
    with METRICS_JSON.open("r", encoding="utf-8") as handle:
        evidence = json.load(handle)
    metrics: dict[str, str] = {}
    with METRICS_CSV.open("r", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            metrics[row["metric"]] = row["value"]
    return evidence, metrics


def metric_value(metrics: dict[str, str], key: str, default: float = 0.0) -> float:
    try:
        return float(metrics.get(key, default))
    except (TypeError, ValueError):
        return default


def format_pct(value: float, decimals: int = 1) -> str:
    return f"{value * 100:.{decimals}f}%"


def format_number(value: float, decimals: int = 0) -> str:
    if decimals == 0:
        return f"{value:,.0f}"
    return f"{value:,.{decimals}f}"


def create_scorecard_chart(metrics: dict[str, str], out_path: Path) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    draw_text(draw, (70, 55), "Executive Evidence Scorecard", 48, COLORS["navy"], True)
    draw_text(
        draw,
        (70, 120),
        "Model quality, runtime reliability, and operational readiness from the collected demo evidence.",
        26,
        COLORS["gray"],
    )

    cards = [
        ("Precision", format_pct(metric_value(metrics, "model_precision"), 2), COLORS["blue"], "Correct anomaly decisions"),
        ("Recall", format_pct(metric_value(metrics, "model_recall"), 2), COLORS["green"], "Anomalies captured"),
        ("F1-score", format_pct(metric_value(metrics, "model_f1_score"), 2), COLORS["orange"], "Balanced model quality"),
        ("API p95", f"{metric_value(metrics, 'api_p95_latency_seconds'):.2f}s", COLORS["navy"], "Observed latency"),
    ]
    x, y = 70, 185
    card_w, card_h, gap = 345, 165, 32
    for i, (label, value, accent, note) in enumerate(cards):
        left = x + i * (card_w + gap)
        draw_rounded_rectangle(draw, (left, y, left + card_w, y + card_h), 22, "FFFFFF", "D9E2EC", 2)
        draw.rectangle((left, y, left + 12, y + card_h), fill=f"#{accent}")
        draw_text(draw, (left + 32, y + 28), label, 26, COLORS["gray"], True)
        draw_text(draw, (left + 32, y + 70), value, 48, accent, True)
        draw_text(draw, (left + 32, y + 128), note, 21, COLORS["gray"])

    bars = [
        ("Production requests", metric_value(metrics, "production_requests"), 400, COLORS["blue"]),
        ("Production anomalies", metric_value(metrics, "production_anomalies"), 400, COLORS["orange"]),
        ("Load-test RPS", metric_value(metrics, "loadtest_rps"), 10, COLORS["green"]),
        ("Load-test p95 ms", metric_value(metrics, "loadtest_p95_ms"), 500, COLORS["navy"]),
        ("Failure rate", metric_value(metrics, "loadtest_failure_rate"), 1, "9B1C1C"),
    ]
    base_y = 430
    draw_text(draw, (70, base_y - 60), "Runtime Evidence", 34, COLORS["navy"], True)
    for idx, (label, value, maximum, accent) in enumerate(bars):
        row_y = base_y + idx * 78
        draw_text(draw, (80, row_y), label, 24, COLORS["black"], True)
        bar_x, bar_y, bar_w, bar_h = 440, row_y + 2, 850, 32
        draw_rounded_rectangle(draw, (bar_x, bar_y, bar_x + bar_w, bar_y + bar_h), 16, "EEF2F6")
        filled = max(0, min(bar_w, int((value / maximum) * bar_w))) if maximum else 0
        if filled > 0:
            draw_rounded_rectangle(draw, (bar_x, bar_y, bar_x + filled, bar_y + bar_h), 16, accent)
        value_text = format_number(value, 2 if value < 10 and value != 0 else 0)
        draw_text(draw, (bar_x + bar_w + 35, row_y - 2), value_text, 25, accent, True)

    draw_text(draw, (70, 838), "Source: reports/demo_evidence/metrics_table.csv", 20, COLORS["gray"])
    img.save(out_path, quality=95)


def create_architecture_chart(out_path: Path) -> None:
    width, height = 1600, 850
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "MLOps Architecture Flow", 48, COLORS["navy"], True)
    draw_text(draw, (70, 113), "From raw server metrics to monitored, auditable anomaly detection.", 26, COLORS["gray"])

    stages = [
        ("Raw server metrics", "CPU, memory, request count,\nerror rate, latency"),
        ("Airflow pipeline", "Validation, preprocessing,\nfeature engineering"),
        ("Isolation Forest", "Model training with\npromotion gates"),
        ("MLflow registry", "Experiment tracking,\nartifacts, production version"),
        ("FastAPI serving", "Authenticated /detect,\n/explain, /drift, /retrain"),
        ("PostgreSQL audit", "Prediction logs,\nrequest IDs, model versions"),
        ("Prometheus/Grafana", "Latency, anomaly rate,\nerrors, drift score"),
        ("Alertmanager bridge", "Warning alerts trigger\nAirflow retraining workflow"),
    ]
    coords = [
        (80, 205), (445, 205), (810, 205), (1175, 205),
        (80, 505), (445, 505), (810, 505), (1175, 505),
    ]
    box_w, box_h = 300, 150
    accent_cycle = [COLORS["blue"], COLORS["orange"], COLORS["green"], COLORS["navy"]]

    for idx, ((title, subtitle), (x, y)) in enumerate(zip(stages, coords)):
        accent = accent_cycle[idx % len(accent_cycle)]
        draw_rounded_rectangle(draw, (x, y, x + box_w, y + box_h), 18, "FFFFFF", "D9E2EC", 2)
        draw.rectangle((x, y, x + box_w, y + 12), fill=f"#{accent}")
        draw_text(draw, (x + 24, y + 32), title, 26, COLORS["navy"], True)
        for line_idx, line in enumerate(subtitle.split("\n")):
            draw_text(draw, (x + 24, y + 74 + line_idx * 30), line, 22, COLORS["gray"])

    def arrow(start: tuple[int, int], end: tuple[int, int], color: str = COLORS["gray"]) -> None:
        draw.line((start, end), fill=f"#{color}", width=4)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        size = 14
        p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
        p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
        draw.polygon([end, p1, p2], fill=f"#{color}")

    def polyline_arrow(points: list[tuple[int, int]], color: str = COLORS["gray"]) -> None:
        for start, end in zip(points, points[1:]):
            draw.line((start, end), fill=f"#{color}", width=4)
        start, end = points[-2], points[-1]
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        size = 14
        p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
        p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
        draw.polygon([end, p1, p2], fill=f"#{color}")

    for i in range(3):
        arrow((coords[i][0] + box_w, coords[i][1] + box_h // 2), (coords[i + 1][0] - 22, coords[i + 1][1] + box_h // 2))
    arrow((coords[3][0] + box_w // 2, coords[3][1] + box_h), (coords[7][0] + box_w // 2, coords[7][1] - 22))
    arrow((coords[4][0] + box_w, coords[4][1] + box_h // 2), (coords[5][0] - 22, coords[5][1] + box_h // 2))
    arrow((coords[5][0] + box_w, coords[5][1] + box_h // 2), (coords[6][0] - 22, coords[6][1] + box_h // 2))
    arrow((coords[6][0] + box_w, coords[6][1] + box_h // 2), (coords[7][0] - 22, coords[7][1] + box_h // 2))
    arrow((coords[1][0] + box_w // 2, coords[1][1] + box_h), (coords[5][0] + box_w // 2, coords[5][1] - 22), COLORS["blue"])
    polyline_arrow(
        [
            (coords[7][0] + box_w // 2, coords[7][1] + box_h),
            (coords[7][0] + box_w // 2, 748),
            (40, 748),
            (40, coords[1][1] + box_h + 24),
            (coords[1][0] - 8, coords[1][1] + box_h + 24),
        ],
        COLORS["orange"],
    )

    draw_text(draw, (70, 798), "Evidence-backed stack: Airflow, MLflow, FastAPI, PostgreSQL, Prometheus, Grafana, Alertmanager.", 20, COLORS["gray"])
    img.save(out_path, quality=95)


def create_explainability_chart(evidence: dict, out_path: Path) -> None:
    features = evidence["api"]["explain"]["body"]["top_features"][:8]
    width, height = 1600, 850
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "Top Drivers of Anomaly Risk", 48, COLORS["navy"], True)
    draw_text(draw, (70, 113), "SHAP tree explainer impact values from the evidence package.", 26, COLORS["gray"])

    max_val = max(item["abs_impact"] for item in features)
    x_label, x_bar, bar_w = 90, 530, 850
    base_y, row_h = 200, 70
    for i, item in enumerate(features):
        y = base_y + i * row_h
        label = item["feature"].replace("_", " ")
        draw_text(draw, (x_label, y), label, 24, COLORS["black"], True)
        draw_rounded_rectangle(draw, (x_bar, y + 5, x_bar + bar_w, y + 38), 16, "EEF2F6")
        fill_w = int((item["abs_impact"] / max_val) * bar_w)
        accent = COLORS["orange"] if i < 3 else COLORS["blue"]
        draw_rounded_rectangle(draw, (x_bar, y + 5, x_bar + fill_w, y + 38), 16, accent)
        draw_text(draw, (x_bar + bar_w + 35, y + 3), f"{item['abs_impact']:.3f}", 24, accent, True)

    draw_text(draw, (70, 795), "Interpretation: higher impact values indicate stronger contribution to anomaly-risk direction.", 20, COLORS["gray"])
    img.save(out_path, quality=95)


def create_fairness_chart(evidence: dict, out_path: Path) -> None:
    groups = evidence["api"]["fairness"]["body"]["group_metrics"]
    width, height = 1600, 820
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw_text(draw, (70, 50), "Operational Fairness by Server Segment", 48, COLORS["navy"], True)
    draw_text(draw, (70, 113), "Anomaly-rate parity is measured by server_id, not human protected attributes.", 26, COLORS["gray"])

    max_rate = 1.0
    base_y, row_h = 210, 82
    x_label, x_bar, bar_w = 110, 470, 850
    for idx, (group, values) in enumerate(groups.items()):
        rate = values["anomaly_rate"]
        req = values["request_count"]
        y = base_y + idx * row_h
        draw_text(draw, (x_label, y), group, 28, COLORS["black"], True)
        draw_text(draw, (x_label, y + 34), f"{req} requests", 20, COLORS["gray"])
        draw_rounded_rectangle(draw, (x_bar, y + 12, x_bar + bar_w, y + 46), 17, "EEF2F6")
        fill_w = int((rate / max_rate) * bar_w)
        accent = COLORS["green"] if rate < 0.35 else COLORS["orange"]
        if fill_w > 0:
            draw_rounded_rectangle(draw, (x_bar, y + 12, x_bar + fill_w, y + 46), 17, accent)
        draw_text(draw, (x_bar + bar_w + 35, y + 10), f"{rate:.3f}", 26, accent, True)

    draw_text(draw, (70, 755), "Max anomaly-rate gap recorded in evidence: 1.0. Review segment-level behavior before production rollout.", 20, COLORS["gray"])
    img.save(out_path, quality=95)


def create_visual_assets(evidence: dict, metrics: dict[str, str]) -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = {
        "scorecard": ASSET_DIR / "scorecard.png",
        "architecture": ASSET_DIR / "architecture_flow.png",
        "explainability": ASSET_DIR / "explainability_drivers.png",
        "fairness": ASSET_DIR / "fairness_segments.png",
    }
    create_scorecard_chart(metrics, assets["scorecard"])
    create_architecture_chart(assets["architecture"])
    create_explainability_chart(evidence, assets["explainability"])
    create_fairness_chart(evidence, assets["fairness"])
    return assets


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb_color(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = rgb_color(color)
    if bold is not None:
        style.font.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11, COLORS["black"])
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    set_style_font(styles["Heading 1"], "Calibri", 16, "2E74B5", True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 2"], "Calibri", 13, "2E74B5", True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True

    set_style_font(styles["Heading 3"], "Calibri", 12, "1F4D78", True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        if style_name in styles:
            styles[style_name].paragraph_format.space_after = Pt(8)
            styles[style_name].paragraph_format.line_spacing = 1.167

    doc.core_properties.title = "final_ddm501 Professional MLOps Project Report"
    doc.core_properties.subject = "Server Log Anomaly Detection MLOps report"
    doc.core_properties.author = "Mai Văn An, Nguyễn Thanh Bình, Phạm Đình Hoàng, Nguyễn Tuấn Dũng"
    doc.core_properties.comments = "Generated from local project evidence and screenshots."


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=COLORS["gray"])
    for instruction in ("PAGE", "NUMPAGES"):
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        paragraph._p.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        paragraph._p.append(instr)
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        paragraph._p.append(fld_sep)
        text = OxmlElement("w:t")
        text.text = "1"
        r = OxmlElement("w:r")
        r.append(text)
        paragraph._p.append(r)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        paragraph._p.append(fld_end)
        if instruction == "PAGE":
            run = paragraph.add_run(" of ")
            set_run_font(run, size=9, color=COLORS["gray"])


def configure_headers_footers(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("final_ddm501 MLOps Report | FPT School of Business & Technology")
    set_run_font(r, size=9, color=COLORS["gray"], bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_p.add_run("final_ddm501 MLOps Report")
    set_run_font(r, size=9, color=COLORS["gray"])

    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Prepared for academic project review | Evidence captured July 1, 2026")
    set_run_font(r, size=8.5, color=COLORS["gray"])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = COLORS["black"], size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def set_table_borders(table, color: str = COLORS["line"], size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_cell_margins(table, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    widths_dxa = [int(width * 1440) for width in widths_in]
    total_dxa = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_dxa):
                continue
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_in: list[float],
    header_fill: str = COLORS["light_gray"],
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_in)
    set_table_borders(table)
    set_table_cell_margins(table)
    header_cells = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], header_fill)
        set_cell_text(header_cells[idx], header, bold=True, color=COLORS["navy"], size=font_size)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value, size=font_size)
    set_table_geometry(table, widths_in)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color=COLORS["gray"], italic=True)


def add_image(doc: Document, path: Path, width: float, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    alt_text = caption or f"Image: {path.stem.replace('_', ' ')}"
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", path.stem)
    if caption:
        add_caption(doc, caption)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        rest = text[len(bold_prefix) :]
        r = p.add_run(rest)
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, title: str, body: str, fill: str = COLORS["soft_blue"], accent: str = COLORS["blue"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="C9D7E6", size="4")
    set_table_cell_margins(table, top=120, bottom=120, start=160, end=160)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=COLORS["black"])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_cover(doc: Document, evidence: dict) -> None:
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    if LOGO_PATH.exists():
        add_image(doc, LOGO_PATH, 5.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Professional Project Report")
    set_run_font(r, size=30, color=COLORS["navy"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Production-Style MLOps for Server Log Anomaly Detection")
    set_run_font(r, size=16, color=COLORS["orange"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    collected = evidence.get("collected_at", "2026-07-01T20:52:31+07:00")
    r = p.add_run(f"Evidence-based English report | Data collected: {collected}")
    set_run_font(r, size=10.5, color=COLORS["gray"], italic=True)

    team_rows = [
        ["Mai Văn An", "Project Lead and MLOps Architecture"],
        ["Nguyễn Thanh Bình", "API Serving, Security, and Runtime Validation"],
        ["Phạm Đình Hoàng", "Model Training, Evaluation, and Explainability"],
        ["Nguyễn Tuấn Dũng", "Observability, CI/CD, and Documentation"],
    ]
    add_table(doc, ["Team Member", "Report Contribution Focus"], team_rows, [2.25, 4.25], header_fill=COLORS["soft_orange"], font_size=10)

    add_callout(
        doc,
        "Report Purpose",
        "This document rewrites the original slide summary into a more complete English report with project context, architecture, evidence metrics, screenshots, tables, and executive recommendations.",
        fill=COLORS["soft_green"],
        accent=COLORS["green"],
    )
    doc.add_page_break()


def add_report_structure(doc: Document) -> None:
    add_heading(doc, "Report Structure", 1)
    rows = [
        ["1", "Executive Summary", "High-level conclusion, evidence highlights, and key recommendation."],
        ["2", "Detailed Slide Summary", "Expanded version of the original summary with metrics and operational meaning."],
        ["3", "Architecture and Implementation", "How the MLOps stack moves from data to served predictions and retraining."],
        ["4", "Model and Runtime Evidence", "Model quality, production predictions, load testing, latency, and auditability."],
        ["5", "Observability and Governance", "Monitoring, explainability, fairness, retraining, and CI/CD readiness."],
    ]
    add_table(doc, ["#", "Section", "Reader Value"], rows, [0.5, 2.1, 3.9], header_fill=COLORS["soft_blue"], font_size=9.4)
    add_para(
        doc,
        "The report is written for an academic evaluation audience that needs both technical credibility and an executive-level explanation of why the project is production-oriented.",
    )


def add_executive_summary(doc: Document, metrics: dict[str, str], evidence: dict, assets: dict[str, Path]) -> None:
    add_heading(doc, "Executive Summary", 1)
    add_callout(
        doc,
        "Executive conclusion",
        "final_ddm501 demonstrates a complete MLOps reference stack for server-log anomaly detection: model training, registry management, authenticated serving, audit logging, observability, alert routing, and retraining workflows are all represented in one runnable Docker Compose environment.",
        fill=COLORS["soft_blue"],
        accent=COLORS["blue"],
    )

    add_para(
        doc,
        "The project goes beyond a standalone machine-learning notebook by packaging the anomaly detector as an operational system. The stack includes Airflow for orchestration, MLflow and MinIO for experiment and artifact management, FastAPI for secured inference endpoints, PostgreSQL for prediction audit trails, Prometheus and Grafana for live telemetry, and Alertmanager for alert-to-retraining automation."
    )
    add_para(
        doc,
        "The collected evidence shows strong model performance and stable demo runtime behavior: precision, recall, and F1-score are each 97.37%; production evidence includes 400 prediction requests with 108 anomalies; API p95 latency is 0.05 seconds; and the Locust smoke load test recorded 7.20 requests per second with zero failures."
    )

    add_image(doc, assets["scorecard"], 6.25, "Figure 1. Evidence scorecard generated from reports/demo_evidence/metrics_table.csv.")

    doc.add_page_break()
    add_heading(doc, "Evidence Highlights", 2)
    rows = [
        ["Model quality", "Precision/Recall/F1 = 97.37%", "The Isolation Forest model performs consistently on the available labeled validation evidence."],
        ["Runtime reliability", "p95 latency = 0.05s; API errors = 0", "The FastAPI service handled demo traffic without recorded runtime errors."],
        ["Operational evidence", "400 production requests; 108 anomalies", "The system records prediction behavior and supports downstream monitoring."],
        ["Load testing", "7.20 RPS; p95 = 210 ms; failure rate = 0", "The local smoke scenario indicates a stable demo serving path."],
        ["Governance", "Explainability + fairness endpoints", "The system exposes interpretable feature impact and segment-level anomaly-rate checks."],
    ]
    add_table(doc, ["Area", "Evidence", "Interpretation"], rows, [1.55, 1.8, 3.15], header_fill=COLORS["light_gray"], font_size=8.8)
    add_heading(doc, "Executive Takeaways", 2)
    takeaways = [
        "The project should be presented as a production-style reference stack because it connects model training, serving, logging, monitoring, alerting, and retraining rather than stopping at offline evaluation.",
        "The strongest evidence points are the consistent model metrics, zero recorded API errors, low p95 latency, and complete observability path through Prometheus and Grafana.",
        "The main limitation is concept drift: without ground-truth production labels, the project can monitor data drift and prediction distribution drift but cannot yet measure post-deployment accuracy drift.",
        "The recommended next step is to harden secrets, persistence, alert thresholds, and larger load tests before treating the stack as a realistic pilot system.",
    ]
    for item in takeaways:
        add_bullet(doc, item)


def add_detailed_slide_summary(doc: Document, metrics: dict[str, str], evidence: dict, assets: dict[str, Path]) -> None:
    add_heading(doc, "Detailed Slide Summary", 1)
    add_para(
        doc,
        "This section rewrites the original slide summary into a fuller English presentation narrative. It can be used as a speaker-note source for a final project presentation or as the condensed executive brief inside the report."
    )

    summary_rows = [
        [
            "Objective",
            "Build a production-style MLOps stack for detecting anomalies in server metrics.",
            "The project targets operational monitoring, not only offline model evaluation. Concept drift with ground-truth labels is intentionally out of scope because production labels are not available.",
        ],
        [
            "Architecture",
            "Raw metrics flow through validation, preprocessing, feature engineering, Isolation Forest training, MLflow tracking, registry promotion, FastAPI serving, PostgreSQL audit logging, and Prometheus/Grafana monitoring.",
            "The architecture shows separation of model lifecycle, serving, observability, and retraining responsibilities.",
        ],
        [
            "Model Result",
            "Precision 0.973684, recall 0.973684, F1-score 0.973684, false positive rate 0.001466, training rows 720.",
            "The model evidence supports a strong demo-quality anomaly detector, with low false-positive behavior on the available validation data.",
        ],
        [
            "Production Runtime",
            "400 requests, 108 anomalies, API p95 latency 0.05s, API error count 0, load-test p95 210 ms, p99 360 ms.",
            "Runtime telemetry suggests the demo stack is responsive and observable under local smoke traffic.",
        ],
        [
            "Observability",
            "Prometheus scrapes FastAPI metrics; Grafana visualizes request rate, p95 latency, anomaly rate, and drift signals; Alertmanager routes warning alerts.",
            "This gives operators the minimum feedback loop required to understand serving health and model behavior.",
        ],
        [
            "Explainability and Fairness",
            "SHAP tree explainer is used for feature impact; operational fairness is measured by server_id segment with max anomaly-rate gap 1.0.",
            "The fairness report is not about human protected classes; it checks whether operational segments behave differently and need investigation.",
        ],
        [
            "Retraining and CI/CD",
            "The /retrain endpoint supports drift checks and forced retraining; Alertmanager can trigger the Airflow retraining DAG; CI runs linting, tests, Docker builds, Trivy scans, image publish, and integration checks.",
            "The system demonstrates a credible path from monitoring signal to retraining and redeployment readiness.",
        ],
    ]
    add_table(doc, ["Slide Topic", "Expanded Summary", "Why It Matters"], summary_rows, [1.25, 2.6, 2.65], header_fill=COLORS["soft_orange"], font_size=8.2)

    add_image(doc, assets["architecture"], 6.25, "Figure 2. High-level architecture flow rewritten from the slide summary.")


def add_architecture_section(doc: Document) -> None:
    add_heading(doc, "Architecture and Implementation", 1)
    add_para(
        doc,
        "The architecture is intentionally production-oriented while remaining runnable on a laptop. Docker Compose coordinates the service layer, while each major production concern is represented by a dedicated component."
    )
    rows = [
        ["Data and feature layer", "Raw server metrics, validation, preprocessing, rolling features, interaction features", "Prepares reliable feature vectors for anomaly detection."],
        ["Model layer", "Isolation Forest, promotion gates, local registry, optional MLflow alias sync", "Controls which model version becomes production-ready."],
        ["Serving layer", "FastAPI endpoints: /detect, /explain, /drift, /fairness, /retrain", "Provides authenticated inference and operational checks."],
        ["Audit layer", "PostgreSQL prediction_logs with request_id, status, prediction, score, version", "Creates traceability for every served prediction."],
        ["Observability layer", "Prometheus metrics, Grafana dashboards, Alertmanager routes", "Tracks health, latency, anomaly rate, drift, and alert state."],
        ["Automation layer", "Airflow DAGs plus Alertmanager bridge", "Connects monitoring signals to retraining workflows."],
        ["Delivery layer", "GitHub Actions, Ruff, pytest, Docker build, Trivy scan, GHCR publish", "Demonstrates repeatable quality gates and image delivery."],
    ]
    add_table(doc, ["Layer", "Key Components", "Purpose"], rows, [1.55, 2.75, 2.2], header_fill=COLORS["soft_blue"], font_size=8.6)

    add_heading(doc, "Service Evidence", 2)
    add_para(
        doc,
        "The evidence package includes screenshots for the running services. These visuals support the claim that the project is not only code-complete but also demonstrable through operational interfaces."
    )
    screenshots = [
        ("fastapi_docs.png", "Figure 3. FastAPI documentation surface for authenticated serving endpoints."),
        ("streamlit_dashboard.png", "Figure 4. Streamlit operations dashboard for demo interaction and monitoring visibility."),
    ]
    for file_name, caption in screenshots:
        path = SCREENSHOT_DIR / file_name
        if path.exists():
            add_image(doc, path, 6.15, caption)


def add_model_runtime_section(doc: Document, metrics: dict[str, str], evidence: dict) -> None:
    add_heading(doc, "Model and Runtime Evidence", 1)
    add_heading(doc, "Model Quality", 2)
    model_rows = [
        ["Precision", "0.973684", "models/latest/metrics.json", "High proportion of predicted anomalies are correct in validation evidence."],
        ["Recall", "0.973684", "models/latest/metrics.json", "High proportion of true anomalies are captured."],
        ["F1-score", "0.973684", "models/latest/metrics.json", "Balanced quality across precision and recall."],
        ["False positive rate", "0.001466", "models/latest/metrics.json", "Low false-positive behavior in the available labeled data."],
        ["Training rows", "720", "models/latest/metrics.json", "Training evidence volume for the demo dataset."],
        ["Production model", evidence["api"]["health"]["body"]["model_version"], "FastAPI /health", "The API loaded the production model from the local registry."],
    ]
    add_table(doc, ["Metric", "Value", "Source", "Interpretation"], model_rows, [1.45, 1.05, 1.7, 2.3], header_fill=COLORS["light_gray"], font_size=8.4)

    mlflow_path = SCREENSHOT_DIR / "mlflow_experiment.png"
    if mlflow_path.exists():
        add_image(doc, mlflow_path, 6.15, "Figure 5. MLflow experiment evidence for model tracking and metrics.")

    doc.add_page_break()
    add_heading(doc, "Production Runtime and Load Test", 2)
    runtime_rows = [
        ["Production requests", format_number(metric_value(metrics, "production_requests")), "PostgreSQL prediction_logs", "Predictions are persisted for audit and analysis."],
        ["Production anomalies", format_number(metric_value(metrics, "production_anomalies")), "PostgreSQL prediction_logs", "Observed anomaly volume across the demo production traffic."],
        ["Rolling anomaly rate", "0.375", "Prometheus", "Operational signal used for monitoring and alerting."],
        ["API p95 latency", f"{metric_value(metrics, 'api_p95_latency_seconds'):.2f}s", "Prometheus", "Serving response time remains low in the demo environment."],
        ["API error count", format_number(metric_value(metrics, "api_error_count")), "Prometheus", "No API errors were recorded in the evidence snapshot."],
        ["Load-test RPS", f"{metric_value(metrics, 'loadtest_rps'):.2f}", "Locust CSV", "Smoke-level throughput for the local demo stack."],
        ["Load-test p95 / p99", "210 ms / 360 ms", "Locust CSV", "Tail latency remains acceptable for demonstration traffic."],
        ["Load-test failure rate", "0", "Locust CSV", "No failures were observed in the captured run."],
    ]
    add_table(doc, ["Runtime Metric", "Value", "Source", "Operational Meaning"], runtime_rows, [1.55, 1.0, 1.55, 2.4], header_fill=COLORS["soft_green"], font_size=8.2)

    grafana_path = SCREENSHOT_DIR / "grafana_dashboard.png"
    if grafana_path.exists():
        add_image(doc, grafana_path, 6.15, "Figure 6. Grafana dashboard evidence for serving and model telemetry.")


def add_observability_governance_section(doc: Document, evidence: dict, assets: dict[str, Path]) -> None:
    add_heading(doc, "Observability, Explainability, and Governance", 1)
    add_heading(doc, "Explainability", 2)
    explain = evidence["api"]["explain"]["body"]
    add_para(
        doc,
        f"The explainability endpoint completed with method `{explain['method']}` for model version `{explain['model_version']}` from `{explain['model_source']}`. For the sample anomalous request, the model returned risk level `{explain['risk_level']}` and anomaly score {explain['anomaly_score']}.",
    )
    feature_rows = [
        [
            item["feature"],
            f"{item['value']}",
            f"{item['abs_impact']:.6f}",
            item["direction"],
        ]
        for item in explain["top_features"][:8]
    ]
    add_table(doc, ["Feature", "Input Value", "Impact", "Direction"], feature_rows, [2.15, 1.2, 1.0, 2.15], header_fill=COLORS["soft_orange"], font_size=8.4)
    add_image(doc, assets["explainability"], 6.2, "Figure 7. Top feature impacts for anomaly-risk explanation.")

    doc.add_page_break()
    add_heading(doc, "Fairness by Operational Segment", 2)
    fairness = evidence["api"]["fairness"]["body"]
    add_para(
        doc,
        "The project does not process human protected attributes. The fairness check therefore measures parity over an operational segment, server_id, so the team can detect whether some servers receive unusually different anomaly rates."
    )
    fairness_rows = [
        [
            group,
            str(values["request_count"]),
            f"{values['anomaly_rate']:.6f}",
            f"{values['high_risk_rate']:.6f}",
            f"{values['avg_anomaly_score']:.6f}",
        ]
        for group, values in fairness["group_metrics"].items()
    ]
    add_table(doc, ["Group", "Requests", "Anomaly Rate", "High-Risk Rate", "Avg. Score"], fairness_rows, [1.35, 1.0, 1.35, 1.3, 1.5], header_fill=COLORS["light_gray"], font_size=8.3)
    add_image(doc, assets["fairness"], 6.15, "Figure 8. Segment-level anomaly-rate comparison from fairness evidence.")

    add_callout(
        doc,
        "Governance note",
        f"The recorded fairness max gap is {fairness['max_anomaly_rate_gap']}. This does not prove unfairness toward people; it flags operational segment imbalance that should be investigated before production rollout.",
        fill=COLORS["soft_orange"],
        accent=COLORS["orange"],
    )


def add_retraining_cicd_section(doc: Document, evidence: dict) -> None:
    add_heading(doc, "Retraining and CI/CD Readiness", 1)
    add_heading(doc, "Retraining Workflow", 2)
    add_para(
        doc,
        "The stack includes both manual and automated retraining paths. The /retrain endpoint can check drift or force a new training job, while the Alertmanager bridge can receive alerts such as DataDriftDetected and HighAnomalyRate and trigger the Airflow retraining DAG through the Airflow REST API."
    )
    drift = evidence["api"]["drift"]["body"]
    retrain = evidence["api"]["retrain_check"]["body"]
    rows = [
        ["Data drift detected", str(drift["data_drift_detected"]), "FastAPI /drift", "Current production features differ from the training reference window."],
        ["Prediction drift detected", str(drift["prediction_drift_detected"]), "FastAPI /drift", "Prediction distribution did not drift above threshold in the captured evidence."],
        ["Max data drift score", f"{drift['max_data_drift_score']:.6f}", "FastAPI /drift", "Largest feature-level drift signal."],
        ["Retraining triggered", str(retrain["retraining_triggered"]), "FastAPI /retrain", retrain["reason"]],
        ["Active alerts", "DataDriftDetected, HighAnomalyRate", "Prometheus ALERTS", "Alertmanager routing is demonstrable in the evidence snapshot."],
    ]
    add_table(doc, ["Check", "Value", "Source", "Meaning"], rows, [1.55, 1.35, 1.45, 2.15], header_fill=COLORS["soft_blue"], font_size=8.4)

    airflow_path = SCREENSHOT_DIR / "airflow_retraining_dag.png"
    if airflow_path.exists():
        add_image(doc, airflow_path, 6.15, "Figure 9. Airflow retraining DAG evidence.")

    add_heading(doc, "CI/CD Evidence", 2)
    github = evidence["github"]
    workflow_runs = github["actions_runs"]["body"].get("workflow_runs", [])
    workflow = workflow_runs[0] if workflow_runs else {}
    jobs = github["latest_run_jobs"]["body"].get("jobs", [])
    job_rows = [
        [job["name"], job["status"], str(job["conclusion"] or "n/a"), job.get("started_at") or "n/a", job.get("completed_at") or "n/a"]
        for job in jobs
    ]
    add_para(
        doc,
        f"The latest recorded CI run in the local evidence package is run #{workflow.get('run_number', 'n/a')} on branch `{workflow.get('head_branch', 'n/a')}`. Completed jobs include Python quality gates, Docker build/scan/publish, and Compose integration testing; the local CD deploy job was queued in the captured evidence."
    )
    add_table(doc, ["Job", "Status", "Conclusion", "Started", "Completed"], job_rows, [2.25, 1.0, 1.0, 1.1, 1.15], header_fill=COLORS["soft_green"], font_size=7.7)


def add_recommendations_conclusion(doc: Document) -> None:
    add_heading(doc, "Recommendations and Conclusion", 1)
    add_heading(doc, "Recommended Next Steps", 2)
    next_steps = [
        "Replace local demo credentials with managed secrets before sharing or deploying outside a classroom environment.",
        "Persist model registry metadata and artifacts outside application containers for production durability.",
        "Add ground-truth label capture when available so concept drift and post-deployment model quality can be measured directly.",
        "Define alert thresholds with operational owners, escalation channels, and retraining approval rules.",
        "Run a larger load test after deployment sizing decisions are made, including concurrent explainability and drift calls.",
    ]
    for step in next_steps:
        add_bullet(doc, step)

    add_heading(doc, "Conclusion", 2)
    add_para(
        doc,
        "The final_ddm501 project is a strong academic demonstration of production-style MLOps. It is especially effective because it connects modeling, serving, monitoring, explainability, fairness checks, drift detection, retraining workflows, CI/CD, and audit logging into one coherent stack."
    )
    add_para(
        doc,
        "The project should be presented as a reference implementation rather than a fully production-certified platform. With stronger secret management, persistent infrastructure, larger-scale performance testing, and ground-truth production labels, it would be ready for a more realistic pilot environment."
    )

    add_heading(doc, "Evidence Sources", 2)
    source_rows = [
        ["Metrics CSV", "reports/demo_evidence/metrics_table.csv", "Core KPI values used in the report."],
        ["Evidence JSON", "reports/demo_evidence/latest_metrics.json", "API, Prometheus, MLflow, Airflow, GitHub, drift, fairness, and explainability evidence."],
        ["Screenshots", "docs/assets/screenshots/*.png", "Visual proof for FastAPI, Streamlit, MLflow, Grafana, and Airflow."],
        ["README", "README.md", "Project scope, architecture, services, CI/CD, and operating instructions."],
        ["Original slide summary", "docs/slide_summary_vi.md", "Vietnamese source summary expanded into this English report."],
    ]
    add_table(doc, ["Source", "Path", "Use"], source_rows, [1.45, 2.55, 2.5], header_fill=COLORS["light_gray"], font_size=8.5)


def build_report() -> None:
    evidence, metrics = load_data()
    assets = create_visual_assets(evidence, metrics)

    doc = Document()
    configure_document(doc)
    configure_headers_footers(doc)

    add_cover(doc, evidence)
    add_report_structure(doc)
    doc.add_page_break()
    add_executive_summary(doc, metrics, evidence, assets)
    doc.add_page_break()
    add_detailed_slide_summary(doc, metrics, evidence, assets)
    doc.add_page_break()
    add_architecture_section(doc)
    doc.add_page_break()
    add_model_runtime_section(doc, metrics, evidence)
    doc.add_page_break()
    add_observability_governance_section(doc, evidence, assets)
    doc.add_page_break()
    add_retraining_cicd_section(doc, evidence)
    doc.add_page_break()
    add_recommendations_conclusion(doc)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
